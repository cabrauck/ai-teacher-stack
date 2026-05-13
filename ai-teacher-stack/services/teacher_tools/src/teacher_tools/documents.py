from __future__ import annotations

from pathlib import Path

from docx import Document

from teacher_tools.models import LessonPlan


def lesson_to_markdown(lesson: LessonPlan) -> str:
    refs = "\n".join(
        f"- {r.subject} / {r.learning_area}: {r.competency} ({r.source.url})"
        for r in lesson.curriculum_references
    ) or "- Kein passender Lehrplanbezug gefunden. Manuell prüfen."

    def bullet(items: list[str]) -> str:
        return "\n".join(f"- {item}" for item in items)

    return f"""# {lesson.title}

## Metadaten

- Fach: {lesson.subject}
- Jahrgangsstufe: {lesson.grade_band}
- Thema: {lesson.topic}
- Dauer: {lesson.duration_minutes} Minuten

## Lehrplanbezug

{refs}

## Lernziele

{bullet(lesson.learning_goals)}

## Materialien

{bullet(lesson.materials)}

## Ablauf

{bullet(lesson.phases)}

## Differenzierung

{bullet(lesson.differentiation)}

## Beobachtung / Sicherung

{bullet(lesson.assessment)}

## Hinweis

{lesson.teacher_review_note}
"""


def export_lesson_docx(lesson: LessonPlan, export_root: Path, filename: str | None = None) -> Path:
    export_root.mkdir(parents=True, exist_ok=True)
    safe_name = filename or f"{lesson.subject}_{lesson.topic}".replace(" ", "_").replace("/", "-")
    if not safe_name.endswith(".docx"):
        safe_name += ".docx"

    path = export_root / safe_name

    doc = Document()
    doc.add_heading(lesson.title, level=1)
    doc.add_paragraph(f"Fach: {lesson.subject}")
    doc.add_paragraph(f"Jahrgangsstufe: {lesson.grade_band}")
    doc.add_paragraph(f"Thema: {lesson.topic}")
    doc.add_paragraph(f"Dauer: {lesson.duration_minutes} Minuten")

    doc.add_heading("Lehrplanbezug", level=2)
    if lesson.curriculum_references:
        for ref in lesson.curriculum_references:
            doc.add_paragraph(
                f"{ref.subject} / {ref.learning_area}: {ref.competency} ({ref.source.url})",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("Kein passender Lehrplanbezug gefunden. Manuell prüfen.")

    sections = [
        ("Lernziele", lesson.learning_goals),
        ("Materialien", lesson.materials),
        ("Ablauf", lesson.phases),
        ("Differenzierung", lesson.differentiation),
        ("Beobachtung / Sicherung", lesson.assessment),
    ]

    for heading, items in sections:
        doc.add_heading(heading, level=2)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Hinweis", level=2)
    doc.add_paragraph(lesson.teacher_review_note)

    doc.save(path)
    return path
